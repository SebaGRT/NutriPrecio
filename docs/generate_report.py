#!/usr/bin/env python3
"""Generate Sprint I report for NutriPrecio in .docx and .pdf formats."""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.lib.colors import HexColor, black, white
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    Image, PageBreak, KeepTogether
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

DOCS_DIR = "/home/sebacc/Documents/UDD/Agil/NutriPrecio/docs"

# Sprint Backlog data from Excel
TASKS = [
    {
        "id": "MV-03",
        "task": "Backend: Crear tabla en BDD para Perfiles de Tienda, vinculada al usuario.",
        "owner": "Vicente Sepulveda",
        "status": "Hecho",
        "estimated": 4,
        "consumed": 2,
    },
    {
        "id": "MV-03",
        "task": "Frontend: Crear formulario para que el vendedor ingrese datos de su pyme.",
        "owner": "Benjamin Buzeta",
        "status": "Hecho",
        "estimated": 5,
        "consumed": 3,
    },
    {
        "id": "MV-03",
        "task": "Backend: Crear endpoint para recibir y guardar la informacion publica de la tienda.",
        "owner": "Matias Ramirez",
        "status": "Hecho",
        "estimated": 5,
        "consumed": 2,
    },
    {
        "id": "MV-52",
        "task": "Backend: Configurar la base de datos para usuarios, y encriptacion de contrasenas.",
        "owner": "Vicente Sepulveda",
        "status": "Hecho",
        "estimated": 4,
        "consumed": 2,
    },
    {
        "id": "MV-52",
        "task": "Frontend: Disenar interfaz del formulario de Registro e Inicio de Sesion.",
        "owner": "Fernando Sepulveda",
        "status": "Hecho",
        "estimated": 5,
        "consumed": 3,
    },
    {
        "id": "MV-52",
        "task": "Backend: Desarrollar logica para validacion de credenciales y tokens de sesion.",
        "owner": "Benjamin Buzeta",
        "status": "Hecho",
        "estimated": 5,
        "consumed": 2,
    },
    {
        "id": "MV-54",
        "task": "Backend: Configurar rutas protegidas en el sistema, para que solo usuarios vendedores puedan ver el panel.",
        "owner": "Vicente Sepulveda",
        "status": "Hecho",
        "estimated": 5,
        "consumed": 2,
    },
    {
        "id": "MV-54",
        "task": "Frontend: Disenar y maquetar el Layout del Dashboard.",
        "owner": "Benjamin Buzeta",
        "status": "Hecho parcialmente",
        "estimated": 8,
        "consumed": 3,
    },
    {
        "id": "MV-54",
        "task": "Frontend: Crear vista de Inicio del Dashboard, mostrando mensaje de bienvenida y estado del perfil.",
        "owner": "Sebastian Herrera",
        "status": "Hecho parcialmente",
        "estimated": 6,
        "consumed": 2,
    },
]

TOTAL_ESTIMATED = sum(t["estimated"] for t in TASKS)
TOTAL_CONSUMED = sum(t["consumed"] for t in TASKS)
TOTAL_REMAINING = TOTAL_ESTIMATED - TOTAL_CONSUMED
COMPLETED = sum(1 for t in TASKS if t["status"] == "Hecho")
PARTIAL = sum(1 for t in TASKS if "parcialmente" in t["status"])


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(doc, text, bold=False, italic=False, size=11,
                          color=None, alignment=None, space_after=6,
                          space_before=0, font_name="Calibri"):
    """Add a paragraph with consistent styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p


def add_mixed_paragraph(doc, parts, alignment=None, space_after=6,
                        space_before=0, font_name="Calibri"):
    """Add paragraph with mixed formatting. parts = [(text, bold, italic, size, color)]."""
    p = doc.add_paragraph()
    for text, bold, italic, size, color in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = font_name
        if color:
            run.font.color.rgb = RGBColor(*color)
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p


def add_heading_styled(doc, text, level=1):
    """Add heading with consistent styling."""
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_heading(level=0)
    run = p.add_run(text)
    run.font.size = Pt(sizes.get(level, 12))
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.bold = True
    pf = p.paragraph_format
    pf.space_before = Pt(12 if level == 1 else 8)
    pf.space_after = Pt(6)
    return p


def add_image_with_caption(doc, image_path, caption, width=5.5):
    """Add image with caption below."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(image_path):
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_p.add_run(caption)
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    caption_p.paragraph_format.space_after = Pt(12)


def create_docx():
    """Create the .docx report."""
    doc = Document()

    # Page setup: A4
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15

    # =====================
    # 1. PORTADA
    # =====================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("NutriPrecio")
    run.bold = True
    run.font.size = Pt(32)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Informe Sprint I")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    subtitle.paragraph_format.space_after = Pt(24)

    # Horizontal line
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_p.add_run("_" * 60)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run.font.size = Pt(10)
    line_p.paragraph_format.space_after = Pt(24)

    info_items = [
        ("Curso:", "Taller de Metodos de Innovacion Agiles"),
        ("Proyecto:", "NutriPrecio - Plataforma de comparacion de precios de alimentos saludables"),
        ("Sprint:", "Sprint I"),
        ("Fecha:", "Mayo 2026"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label + " ")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run = p.add_run(value)
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    team_p = doc.add_paragraph()
    team_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = team_p.add_run("Equipo de Desarrollo")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    team_p.paragraph_format.space_after = Pt(8)

    team_members = [
        ("Sebastian Herrera", "Scrum Master"),
        ("Ignacio Herrera", "Product Owner"),
        ("Vicente Sepulveda", "Developer"),
        ("Matias Ramirez", "Developer"),
        ("Benjamin Buzeta", "Developer"),
        ("Fernando Sepulveda", "Developer"),
    ]
    for name, role in team_members:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{name} ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run = p.add_run(f"- {role}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(3)

    # Page break after cover
    doc.add_page_break()

    # =====================
    # 2. INTRODUCCION
    # =====================
    add_heading_styled(doc, "1. Introduccion", level=1)

    add_styled_paragraph(doc,
        "El presente informe documenta el desarrollo y los resultados obtenidos durante el Sprint I del proyecto NutriPrecio, una plataforma web de comparacion de precios de alimentos saludables orientada al mercado chileno. El sistema permite a los consumidores buscar, comparar y encontrar los mejores precios en supermercados y tiendas locales de Chile.",
        size=11, space_after=10)

    add_heading_styled(doc, "1.1 Objetivo del Sprint I", level=2)

    add_styled_paragraph(doc,
        "El objetivo principal del Sprint I fue implementar las funcionalidades base de la plataforma NutriPrecio, centradas en tres historias de usuario fundamentales:",
        size=11, space_after=6)

    objectives = [
        "MV-03: Registro de tienda para vendedores independientes, permitiendo que pymes y emprendedores registren su negocio en la plataforma.",
        "MV-52: Sistema de login y registro con autenticacion por tokens, habilitando la creacion de cuentas y el acceso seguro a funcionalidades protegidas.",
        "MV-54: Panel de control privado para vendedores, donde puedan gestionar la informacion de su tienda en un solo lugar.",
    ]
    for obj in objectives:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(obj)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(4)

    add_heading_styled(doc, "1.2 Alcance del Incremento", level=2)

    add_styled_paragraph(doc,
        "El alcance del Sprint I comprendio el desarrollo del sistema de autenticacion completo (backend y frontend), el modelo de datos para tiendas con sus endpoints API, el formulario de registro de tienda en el frontend, y la vista inicial del dashboard del vendedor. Se planificaron 9 tareas con un total de 47 horas estimadas, distribuidas entre los 4 desarrolladores del equipo.",
        size=11, space_after=10)

    # =====================
    # 3. METODOLOGIA
    # =====================
    add_heading_styled(doc, "2. Metodologia y Enfoque Tecnico", level=1)

    add_heading_styled(doc, "2.1 Framework Scrum", level=2)

    add_styled_paragraph(doc,
        "El equipo adopto el framework Scrum para la gestion del Sprint I. La planificacion se realizo mediante la seleccion de items del Product Backlog priorizados por el Product Owner, los cuales se desglosaron en tareas tecnicas asignadas a cada miembro del equipo. El Sprint Backlog se registro en una plantilla oficial de seguimiento, con estimacion de horas por tarea y registro del consumo diario.",
        size=11, space_after=10)

    add_heading_styled(doc, "2.2 Herramientas Utilizadas", level=2)

    tools = [
        ("Backend:", "Django 6.0 con Django REST Framework para la API REST, SQLite para desarrollo, django-cors-headers para CORS, y Django Token Authentication para seguridad."),
        ("Frontend:", "Angular 19 con componentes standalone, Angular Material para la interfaz de usuario, SCSS para estilos, y RxJS para manejo de estado mediante servicios con Observables."),
        ("Control de versiones:", "Git para gestion de codigo, con ramas por feature y commits en modo imperativo."),
        ("Colaboracion:", "Plantilla de Sprint Backlog en Excel para seguimiento de tareas y horas."),
    ]
    for label, desc in tools:
        p = doc.add_paragraph()
        run = p.add_run(label + " ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run = p.add_run(desc)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(4)

    add_heading_styled(doc, "2.3 Roles del Equipo", level=2)

    roles_table = doc.add_table(rows=7, cols=2)
    roles_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    roles_table.style = "Table Grid"

    headers = ["Rol", "Nombre"]
    for i, header in enumerate(headers):
        cell = roles_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1F3864")

    roles_data = [
        ("Scrum Master", "Sebastian Herrera"),
        ("Product Owner", "Ignacio Herrera"),
        ("Developer", "Vicente Sepulveda"),
        ("Developer", "Matias Ramirez"),
        ("Developer", "Benjamin Buzeta"),
        ("Developer", "Fernando Sepulveda"),
    ]
    for idx, (rol, nombre) in enumerate(roles_data):
        row = roles_table.rows[idx + 1]
        row.cells[0].text = rol
        row.cells[1].text = nombre
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
            if idx % 2 == 0:
                set_cell_shading(cell, "D6E4F0")

    doc.add_paragraph()

    add_styled_paragraph(doc,
        f"Resumen: {COMPLETED} de 9 tareas completadas (78%), {PARTIAL} tareas parciales (22%). "
        f"Horas estimadas: {TOTAL_ESTIMATED}. Horas consumidas: {TOTAL_CONSUMED}. "
        f"Horas restantes: {TOTAL_REMAINING}.",
        bold=True, size=10, space_after=12)

    # =====================
    # 5. DIAGRAMA DE CASOS DE USO
    # =====================
    add_heading_styled(doc, "4. Diagrama de Casos de Uso", level=1)

    add_image_with_caption(doc,
        os.path.join(DOCS_DIR, "diagrama-casos-de-uso.png"),
        "Figura 1: Diagrama de Casos de Uso del sistema NutriPrecio")

    add_heading_styled(doc, "4.1 Actores del Sistema", level=2)

    actors = [
        ("Usuario No Autenticado:", "Puede navegar por la plataforma, buscar productos, ver categorias y comparar precios sin necesidad de registrarse. Tiene acceso de solo lectura a la informacion publica."),
        ("Usuario Registrado (Comprador):", "Puede crear una cuenta, iniciar sesion, y acceder a funcionalidades personalizadas como listas de favoritos y historial de busquedas."),
        ("Vendedor (Seller):", "Usuario registrado con rol de vendedor. Puede registrar su tienda, gestionar su perfil, y acceder al panel de control privado para administrar su negocio en la plataforma."),
        ("Sistema de Autenticacion:", "Componente interno que gestiona el registro de usuarios, validacion de credenciales, generacion y verificacion de tokens de sesion."),
    ]
    for label, desc in actors:
        p = doc.add_paragraph()
        run = p.add_run(label + " ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run = p.add_run(desc)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(4)

    add_heading_styled(doc, "4.2 Casos de Uso Principales", level=2)

    use_cases = [
        "Registrarse: El usuario crea una cuenta con email y contrasena encriptada.",
        "Iniciar Sesion: El usuario ingresa sus credenciales y recibe un token de autenticacion.",
        "Buscar Productos: El usuario busca productos por nombre, categoria o marca.",
        "Comparar Precios: El usuario compara precios de un producto en diferentes tiendas.",
        "Registrar Tienda: El vendedor registra su tienda con nombre, logo y sitio web.",
        "Acceder al Dashboard: El vendedor accede a su panel de control privado.",
    ]
    for uc in use_cases:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(uc)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(3)

    # =====================
    # 6. DIAGRAMA DE CLASES
    # =====================
    add_heading_styled(doc, "5. Diagrama de Clases", level=1)

    add_image_with_caption(doc,
        os.path.join(DOCS_DIR, "diagrama-clases.png"),
        "Figura 2: Diagrama de Clases del sistema NutriPrecio")

    add_heading_styled(doc, "5.1 Estructura del Sistema", level=2)

    add_styled_paragraph(doc,
        "El diagrama de clases muestra la arquitectura del sistema NutriPrecio, organizada en capas que separan la interfaz de usuario, la logica de negocio y el acceso a datos.",
        size=11, space_after=8)

    add_heading_styled(doc, "5.2 Clases Principales", level=2)

    classes_desc = [
        ("User (Modelo):", "Extiende AbstractUser de Django. Contiene email unico, fecha de creacion y actualizacion. Relacion uno-a-muchos con Store a traves del campo owner."),
        ("Store (Modelo):", "Representa una tienda o pyme. Incluye nombre, slug unico, logo, sitio web, estado activo, y vinculacion al usuario vendedor (owner)."),
        ("Product (Modelo):", "Catalogo de productos con nombre, marca, categoria, unidad de medida, imagen, descripcion y codigo de barras."),
        ("Category (Modelo):", "Categorias jerarquicas de productos. Puede tener una categoria padre (self-referencia)."),
        ("Price (Modelo):", "Registro de precios que vincula un producto con una tienda. Incluye precio actual, precio original, porcentaje de descuento, URL, disponibilidad y fecha de registro."),
        ("AuthService:", "Servicio Angular que gestiona login, registro y logout. Almacena el token de autenticacion en localStorage."),
        ("ApiService:", "Servicio Angular base para comunicaciones HTTP. Provee metodos get, post, put, delete tipados."),
        ("AuthGuard:", "Guard de rutas Angular que verifica autenticacion y rol de vendedor antes de permitir acceso a rutas protegidas."),
    ]
    for label, desc in classes_desc:
        p = doc.add_paragraph()
        run = p.add_run(label + " ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run = p.add_run(desc)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(4)

    add_heading_styled(doc, "5.3 Relaciones", level=2)

    relations = [
        "User (1) -- (N) Store: Un vendedor puede tener una o mas tiendas.",
        "Category (1) -- (N) Product: Una categoria contiene multiples productos.",
        "Category (1) -- (0..1) Category: Relacion jerarquica padre-hijo entre categorias.",
        "Product (1) -- (N) Price: Un producto tiene multiples registros de precio.",
        "Store (1) -- (N) Price: Una tienda tiene multiples registros de precio.",
    ]
    for rel in relations:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(rel)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(3)

    # =====================
    # 7. DIAGRAMA DE ACTIVIDADES
    # =====================
    add_heading_styled(doc, "6. Diagrama de Actividades", level=1)

    add_image_with_caption(doc,
        os.path.join(DOCS_DIR, "diagrama-actividades.png"),
        "Figura 3: Diagrama de Actividades con los 3 flujos principales")

    add_heading_styled(doc, "6.1 Flujo MV-52: Login y Registro", level=2)

    add_styled_paragraph(doc,
        "Este flujo describe el proceso de autenticacion del usuario. Comienza cuando el usuario accede al formulario de login o registro en el frontend. El sistema valida los campos ingresados, envia la solicitud al backend, que verifica las credenciales contra la base de datos. Si son validas, genera un token de autenticacion y lo retorna al frontend, que lo almacena para peticiones futuras. En caso de registro, se crea el usuario con contrasena encriptada mediante create_user() de Django.",
        size=11, space_after=10)

    add_heading_styled(doc, "6.2 Flujo MV-03: Registro de Tienda", level=2)

    add_styled_paragraph(doc,
        "El flujo de registro de tienda inicia cuando un vendedor autenticado accede al formulario desde el dashboard. El frontend construye un objeto FormData con los datos de la tienda (nombre, logo, sitio web) y lo envia al endpoint POST /api/stores/. El backend valida que el usuario tenga rol de vendedor, crea la instancia de Store vinculada al usuario como owner, y retorna la tienda creada. El frontend navega de vuelta al dashboard con confirmacion visual.",
        size=11, space_after=10)

    add_heading_styled(doc, "6.3 Flujo MV-54: Dashboard del Vendedor", level=2)

    add_styled_paragraph(doc,
        "Cuando un vendedor accede a /dashboard, el authGuard verifica que este autenticado y que tenga el rol is_seller. Si ambas condiciones se cumplen, se carga la vista del dashboard con el mensaje de bienvenida personalizado, el perfil del usuario con avatar de iniciales, indicadores de estado (Cuenta Activa, chip de Vendedor), y acciones rapidas que incluyen navegacion al formulario de registro de tienda. Si el usuario no es vendedor, es redirigido al home con un mensaje de aviso.",
        size=11, space_after=10)

    # =====================
    # 8. SPRINT REVIEW
    # =====================
    add_heading_styled(doc, "7. Sprint Review", level=1)

    add_heading_styled(doc, "7.1 Funcionalidades Entregadas", level=2)

    delivered = [
        ("MV-52: Login y Registro (COMPLETO)",
         "Registro de usuarios con encriptacion de contrasenas mediante create_user() de Django. Inicio de sesion con generacion y retorno de tokens de autenticacion. Interceptor HTTP en el frontend (authInterceptor) que incluye el token automaticamente en cada request. Guard de rutas protegidas (authGuard) que verifica existencia del token y redirige al login si no hay sesion activa. Formularios de registro e inicio de sesion con validacion de campos en el frontend."),
        ("MV-03: Registro de Tienda (COMPLETO)",
         "Modelo Store con campo owner (ForeignKey a User) que vincula cada tienda al usuario vendedor. Serializer y ViewSet con permisos de autenticacion para creacion y edicion, acceso publico de solo lectura para listados. Endpoint CRUD funcional en /api/stores/. Formulario de registro de tienda en el frontend (store-form.component) con campos de nombre, logo y sitio web, integrado al dashboard mediante accion rapida."),
        ("MV-54: Dashboard del Vendedor (COMPLETO)",
         "Vista de inicio del dashboard con mensaje de bienvenida personalizado usando el nombre del usuario. Perfil del usuario con avatar de iniciales, correo electronico y nombre completo. Indicadores de estado visual (Cuenta Activa, chip de Vendedor para usuarios is_seller). Acciones rapidas con navegacion al formulario de registro de tienda. Validacion de rol vendedor en el authGuard: usuarios no vendedores son redirigidos al home con mensaje de aviso."),
    ]
    for title, desc in delivered:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
        p.paragraph_format.space_after = Pt(3)
        p = doc.add_paragraph()
        run = p.add_run(desc)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(8)

    add_heading_styled(doc, "7.2 Funcionalidades NO Entregadas / Parciales", level=2)

    partial_items = [
        ("Dashboard con sidebar de navegacion:",
         "El dashboard actual es una vista de una sola pagina con secciones de perfil, estadisticas y acciones rapidas. No se implemento el layout con barra lateral de navegacion que se muestra en el diagrama de actividades."),
        ("Estado real del perfil de tienda:",
         "El dashboard muestra tarjetas de estadisticas con valor vacio y texto Proximamente porque no hay integracion activa con la API para consultar datos reales de la tienda del vendedor."),
        ("Gestion de productos desde el dashboard:",
         "El boton de Configurar Perfil en las acciones rapidas esta deshabilitado con etiqueta Proximamente. La gestion de productos (CRUD) no forma parte del alcance del Sprint I."),
    ]
    for title, desc in partial_items:
        p = doc.add_paragraph()
        run = p.add_run(title + " ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0xBF, 0x8F, 0x00)
        run = p.add_run(desc)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(6)

    add_heading_styled(doc, "7.3 Metricas del Sprint", level=2)

    metrics_table = doc.add_table(rows=6, cols=2)
    metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    metrics_table.style = "Table Grid"

    metrics_headers = ["Indicador", "Valor"]
    for i, header in enumerate(metrics_headers):
        cell = metrics_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1F3864")

    metrics_data = [
        ("Tareas planificadas", "9"),
        ("Tareas completadas", "7 de 9 (78%)"),
        ("Tareas parciales", "2 de 9 (22%)"),
        ("Horas estimadas", "47 horas"),
        ("Horas consumidas", f"{TOTAL_CONSUMED} horas"),
    ]
    for idx, (ind, val) in enumerate(metrics_data):
        row = metrics_table.rows[idx + 1]
        row.cells[0].text = ind
        row.cells[1].text = val
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
            if idx % 2 == 0:
                set_cell_shading(cell, "D6E4F0")

    doc.add_paragraph()

    add_heading_styled(doc, "7.4 Feedback Obtenido", level=2)

    add_styled_paragraph(doc,
        "No se obtuvo feedback formal de stakeholders durante el Sprint I. Las revisiones se realizaron internamente entre miembros del equipo durante el desarrollo.",
        size=11, space_after=10)

    # =====================
    # 9. SPRINT RETROSPECTIVE
    # =====================
    add_heading_styled(doc, "8. Sprint Retrospective", level=1)

    add_heading_styled(doc, "8.1 Que Salio Bien", level=2)

    good_items = [
        "Distribucion clara de tareas: Cada integrante del equipo tuvo responsabilidades bien definidas desde el inicio del Sprint. Vicente se enfoco en el backend (modelos, permisos), Benjamin en el frontend (formularios, dashboard), Matias en endpoints API, Fernando en interfaces de autenticacion, y Sebastian en la vista del dashboard.",
        "Stack tecnologico apropiado: Django + Django REST Framework para el backend y Angular 19 para el frontend demostraron ser una combinacion solida. La separacion de concerns entre apps de Django y componentes standalone de Angular facilito el desarrollo paralelo.",
        "Autenticacion implementada correctamente desde el inicio: El sistema de tokens de Django REST Framework se configuro correctamente desde el primer dia, incluyendo el interceptor HTTP en Angular que automaticamente incluye el token en cada request.",
        "Estructura del proyecto organizada: Las apps de Django separadas por dominio (users, stores, products, prices, categories) y la estructura de carpetas del frontend (core/services, features, shared/components) permitieron que multiples desarrolladores trabajaran sin conflictos.",
    ]
    for item in good_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(4)

    add_heading_styled(doc, "8.2 Que Se Puede Mejorar", level=2)

    improve_items = [
        "Falta de seguimiento diario y check-ins frecuentes: El equipo no tuvo standups diarios formales, lo que resulto en tareas que se estancaron sin que nadie lo notara a tiempo.",
        "Campo owner de Store no se incluyo desde el inicio: El modelo Store se creo sin el campo owner (ForeignKey a User), lo que contradecia el diagrama de clases. Esto requirio una migracion adicional y ajustes en el serializer y viewset.",
        "No se valido el rol is_seller en el dashboard inicialmente: El authGuard solo verificaba si el usuario estaba autenticado, no si era vendedor. Esto permitio que usuarios compradores accedieran al dashboard.",
        "Sin datos de prueba hasta el final: La base de datos estuvo vacia durante todo el Sprint. No se creo seed data hasta el ultimo momento, lo que imposibilito hacer demos intermedias y pruebas reales de integracion.",
        "Formulario de registro de tienda subestimado: La tarea de crear el formulario frontend tomo mas tiempo del estimado porque requirio integracion con el backend (FormData para file upload), manejo de errores, y navegacion.",
    ]
    for item in improve_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(4)

    add_heading_styled(doc, "8.3 Acciones para el Proximo Sprint", level=2)

    actions = [
        "Implementar daily standups: Reuniones diarias de 15 minutos para reportar progreso, bloqueos, y plan del dia.",
        "Crear seed data al inicio del Sprint: Tener datos de prueba desde el dia 1 para permitir demos y pruebas continuas.",
        "Revisar modelos contra diagramas UML antes de codificar: Validar que los modelos Django reflejen exactamente lo disenado en los diagramas de clases.",
        "Definir criterios de aceptacion mas claros por tarea: Cada tarea debe tener criterios de aceptacion especificos y verificables antes de comenzar.",
        "Estimar con mas margen para tareas de integracion: Las tareas que involucran frontend + backend deben tener estimaciones mas conservadoras.",
    ]
    for i, action in enumerate(actions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {action}")
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(4)

    # =====================
    # 10. CONCLUSIONES
    # =====================
    add_heading_styled(doc, "9. Conclusiones", level=1)

    add_heading_styled(doc, "9.1 Lecciones Aprendidas", level=2)

    lessons = [
        "La comunicacion constante entre el diseno UML y la implementacion es fundamental. Las discrepancias entre diagramas y codigo generan retrabajo y retrasos.",
        "Las tareas de integracion frontend-backend requieren mas tiempo del estimado inicialmente. Es necesario considerar la complejidad de la comunicacion entre capas al estimar.",
        "Los datos de prueba son esenciales desde el inicio del Sprint. Sin seed data, no es posible realizar demos intermedias ni validar la integracion de componentes.",
        "La validacion de roles y permisos debe ser parte de los criterios de aceptacion desde el inicio, no un agregado posterior.",
        "El stack Django + Angular demostro ser una combinacion efectiva para desarrollo paralelo, siempre que los contratos de API esten bien definidos.",
    ]
    for lesson in lessons:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(lesson)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(4)

    add_heading_styled(doc, "9.2 Plan para Sprint II", level=2)

    add_styled_paragraph(doc,
        "Para el Sprint II, el equipo se enfocara en completar las funcionalidades parciales del Sprint I y avanzar en nuevas caracteristicas de la plataforma:",
        size=11, space_after=6)

    sprint2_items = [
        "Completar el layout del dashboard con sidebar de navegacion lateral.",
        "Integrar el dashboard con la API para mostrar datos reales de la tienda (productos publicados, estado de la tienda).",
        "Implementar la gestion de productos (CRUD) desde el dashboard del vendedor.",
        "Desarrollar la funcionalidad de busqueda y comparacion de precios para compradores.",
        "Implementar daily standups y mejorar el seguimiento de tareas.",
        "Crear y mantener seed data actualizada para pruebas y demos.",
    ]
    for item in sprint2_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(4)

    add_styled_paragraph(doc,
        "El Sprint I sento las bases tecnicas y organizacionales del proyecto NutriPrecio. A pesar de las tareas parciales, el equipo logro entregar un incremento funcional con autenticacion completa, registro de tiendas y dashboard del vendedor. Las lecciones aprendidas serviran para mejorar la planificacion y ejecucion del Sprint II.",
        size=11, space_after=12, space_before=12)

    # Add footer with page numbers
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("NutriPrecio - Informe Sprint I  |  ")
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run = p.add_run("Pagina ")
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        # Add page number field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        p.runs[-1]._r.addnext(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        fldChar1.addnext(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        instrText.addnext(fldChar2)

    # Save
    output_path = os.path.join(DOCS_DIR, "Informe-Sprint-I-NutriPrecio.docx")
    doc.save(output_path)
    print(f"DOCX saved: {output_path}")
    return output_path


def create_pdf(docx_path):
    """Create PDF version of the report using reportlab."""
    output_path = os.path.join(DOCS_DIR, "Informe-Sprint-I-NutriPrecio.pdf")

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    styles.add(ParagraphStyle(
        name="ReportTitle",
        fontName="Helvetica-Bold",
        fontSize=28,
        textColor=HexColor("#1F3864"),
        alignment=TA_CENTER,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="ReportSubtitle",
        fontName="Helvetica-Bold",
        fontSize=18,
        textColor=HexColor("#4472C4"),
        alignment=TA_CENTER,
        spaceAfter=20,
    ))
    styles.add(ParagraphStyle(
        name="ReportHeading1",
        fontName="Helvetica-Bold",
        fontSize=16,
        textColor=HexColor("#1F3864"),
        spaceBefore=16,
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="ReportHeading2",
        fontName="Helvetica-Bold",
        fontSize=13,
        textColor=HexColor("#1F3864"),
        spaceBefore=10,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="ReportBody",
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="ReportBullet",
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        leftIndent=20,
        spaceAfter=4,
        bulletIndent=8,
    ))
    styles.add(ParagraphStyle(
        name="ReportCenter",
        fontName="Helvetica",
        fontSize=10,
        alignment=TA_CENTER,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="ReportCaption",
        fontName="Helvetica-Oblique",
        fontSize=8,
        textColor=HexColor("#666666"),
        alignment=TA_CENTER,
        spaceAfter=12,
    ))
    styles.add(ParagraphStyle(
        name="ReportFooter",
        fontName="Helvetica",
        fontSize=7,
        textColor=HexColor("#999999"),
        alignment=TA_CENTER,
    ))

    story = []

    # Helper functions
    def add_title(text, style="ReportTitle"):
        story.append(Paragraph(text, styles[style]))

    def add_heading(text, level=1):
        style = "ReportHeading1" if level == 1 else "ReportHeading2"
        story.append(Paragraph(text, styles[style]))

    def add_body(text):
        story.append(Paragraph(text, styles["ReportBody"]))

    def add_bullet(text):
        story.append(Paragraph(f"\u2022 {text}", styles["ReportBullet"]))

    def add_spacer(height=12):
        story.append(Spacer(1, height))

    def add_table(headers, data, col_widths=None):
        table_data = [headers] + data
        table = Table(table_data, colWidths=col_widths)
        table_style = [
            ("BACKGROUND", (0, 0), (-1, 0), HexColor("#1F3864")),
            ("TEXTCOLOR", (0, 0), (-1, 0), white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 1), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ]
        # Alternate row colors
        for i in range(1, len(data) + 1):
            if i % 2 == 0:
                table_style.append(("BACKGROUND", (0, i), (-1, i), HexColor("#F2F2F2")))
        table.setStyle(TableStyle(table_style))
        story.append(table)
        add_spacer(6)

    def add_image_with_caption(image_path, caption, width=12 * cm, max_height=22 * cm):
        if os.path.exists(image_path):
            img = Image(image_path, width=width)
            # Cap height to fit on page
            if img.drawHeight > max_height:
                img.drawHeight = max_height
                # Recalculate width to maintain aspect ratio
                from reportlab.lib.utils import ImageReader
                reader = ImageReader(image_path)
                orig_w, orig_h = reader.getSize()
                scale = max_height / orig_h
                img.drawWidth = orig_w * scale
            img.hAlign = "CENTER"
            story.append(img)
        story.append(Paragraph(caption, styles["ReportCaption"]))

    # =====================
    # 1. PORTADA
    # =====================
    for _ in range(6):
        add_spacer(20)

    add_title("NutriPrecio")
    add_title("Informe Sprint I", "ReportSubtitle")

    add_spacer(12)
    story.append(Paragraph("_" * 50, styles["ReportCenter"]))
    add_spacer(12)

    info_lines = [
        ("Curso:", "Taller de Metodos de Innovacion Agiles"),
        ("Proyecto:", "NutriPrecio - Plataforma de comparacion de precios"),
        ("Sprint:", "Sprint I"),
        ("Fecha:", "Mayo 2026"),
    ]
    for label, value in info_lines:
        p = Paragraph(f'<b>{label}</b> {value}', styles["ReportCenter"])
        story.append(p)

    add_spacer(16)
    story.append(Paragraph('<b>Equipo de Desarrollo</b>', styles["ReportCenter"]))
    add_spacer(6)

    team = [
        "Sebastian Herrera - Scrum Master",
        "Ignacio Herrera - Product Owner",
        "Vicente Sepulveda - Developer",
        "Matias Ramirez - Developer",
        "Benjamin Buzeta - Developer",
        "Fernando Sepulveda - Developer",
    ]
    for member in team:
        story.append(Paragraph(member, styles["ReportCenter"]))

    story.append(PageBreak())

    # =====================
    # 2. INTRODUCCION
    # =====================
    add_heading("1. Introduccion")

    add_body(
        "El presente informe documenta el desarrollo y los resultados obtenidos durante el Sprint I "
        "del proyecto NutriPrecio, una plataforma web de comparacion de precios de alimentos saludables "
        "orientada al mercado chileno. El sistema permite a los consumidores buscar, comparar y encontrar "
        "los mejores precios en supermercados y tiendas locales de Chile."
    )

    add_heading("1.1 Objetivo del Sprint I", level=2)

    add_body(
        "El objetivo principal del Sprint I fue implementar las funcionalidades base de la plataforma "
        "NutriPrecio, centradas en tres historias de usuario fundamentales:"
    )

    add_bullet("MV-03: Registro de tienda para vendedores independientes.")
    add_bullet("MV-52: Sistema de login y registro con autenticacion por tokens.")
    add_bullet("MV-54: Panel de control privado para vendedores.")

    add_heading("1.2 Alcance del Incremento", level=2)

    add_body(
        "El alcance del Sprint I comprendio el desarrollo del sistema de autenticacion completo "
        "(backend y frontend), el modelo de datos para tiendas con sus endpoints API, el formulario "
        "de registro de tienda en el frontend, y la vista inicial del dashboard del vendedor. Se "
        "planificaron 9 tareas con un total de 47 horas estimadas, distribuidas entre los 4 "
        "desarrolladores del equipo."
    )

    # =====================
    # 3. METODOLOGIA
    # =====================
    add_heading("2. Metodologia y Enfoque Tecnico")

    add_heading("2.1 Framework Scrum", level=2)

    add_body(
        "El equipo adopto el framework Scrum para la gestion del Sprint I. La planificacion se realizo "
        "mediante la seleccion de items del Product Backlog priorizados por el Product Owner, los cuales "
        "se desglosaron en tareas tecnicas asignadas a cada miembro del equipo."
    )

    add_heading("2.2 Herramientas Utilizadas", level=2)

    tools = [
        ("Backend:", "Django 6.0 con Django REST Framework, SQLite, django-cors-headers, Token Authentication."),
        ("Frontend:", "Angular 19 con componentes standalone, Angular Material, SCSS, RxJS."),
        ("Control de versiones:", "Git con ramas por feature y commits en modo imperativo."),
        ("Colaboracion:", "Plantilla de Sprint Backlog en Excel para seguimiento de tareas."),
    ]
    for label, desc in tools:
        p = Paragraph(f'<b>{label}</b> {desc}', styles["ReportBody"])
        story.append(p)

    add_heading("2.3 Roles del Equipo", level=2)

    roles_headers = ["Rol", "Nombre"]
    roles_data = [
        ["Scrum Master", "Sebastian Herrera"],
        ["Product Owner", "Ignacio Herrera"],
        ["Developer", "Vicente Sepulveda"],
        ["Developer", "Matias Ramirez"],
        ["Developer", "Benjamin Buzeta"],
        ["Developer", "Fernando Sepulveda"],
    ]
    add_table(roles_headers, roles_data, col_widths=[3*cm, 10*cm])

    # =====================
    # 4. SPRINT BACKLOG
    # =====================
    add_heading("3. Sprint Backlog")

    add_body(
        "A continuacion se presenta el detalle de las 9 tareas planificadas para el Sprint I:"
    )

    backlog_headers = ["ID", "Tarea", "Responsable", "Hrs Est.", "Estado"]
    backlog_data = []
    for t in TASKS:
        task_short = t["task"][:55] + "..." if len(t["task"]) > 55 else t["task"]
        backlog_data.append([
            t["id"], task_short, t["owner"], str(t["estimated"]), t["status"]
        ])
    backlog_data.append([
        "", "TOTAL", "", str(TOTAL_ESTIMATED), f"{COMPLETED} Hecho, {PARTIAL} Parcial"
    ])
    add_table(backlog_headers, backlog_data,
              col_widths=[1.5*cm, 7*cm, 3*cm, 1.5*cm, 3*cm])

    add_body(
        f"<b>Resumen:</b> {COMPLETED} de 9 tareas completadas (78%), {PARTIAL} tareas parciales (22%). "
        f"Horas estimadas: {TOTAL_ESTIMATED}. Horas consumidas: {TOTAL_CONSUMED}."
    )

    # =====================
    # 5. DIAGRAMA DE CASOS DE USO
    # =====================
    add_heading("4. Diagrama de Casos de Uso")

    add_image_with_caption(
        os.path.join(DOCS_DIR, "diagrama-casos-de-uso.png"),
        "Figura 1: Diagrama de Casos de Uso del sistema NutriPrecio"
    )

    add_heading("4.1 Actores del Sistema", level=2)

    actors = [
        ("Usuario No Autenticado:", "Navega, busca productos y compara precios sin registro."),
        ("Usuario Registrado (Comprador):", "Crea cuenta, inicia sesion, accede a funcionalidades personalizadas."),
        ("Vendedor (Seller):", "Registra su tienda, gestiona su perfil, accede al dashboard privado."),
        ("Sistema de Autenticacion:", "Gestiona registro, validacion de credenciales y tokens de sesion."),
    ]
    for label, desc in actors:
        p = Paragraph(f'<b>{label}</b> {desc}', styles["ReportBody"])
        story.append(p)

    add_heading("4.2 Casos de Uso Principales", level=2)

    for uc in ["Registrarse", "Iniciar Sesion", "Buscar Productos", "Comparar Precios",
               "Registrar Tienda", "Acceder al Dashboard"]:
        add_bullet(uc)

    # =====================
    # 6. DIAGRAMA DE CLASES
    # =====================
    add_heading("5. Diagrama de Clases")

    add_image_with_caption(
        os.path.join(DOCS_DIR, "diagrama-clases.png"),
        "Figura 2: Diagrama de Clases del sistema NutriPrecio"
    )

    add_heading("5.1 Estructura del Sistema", level=2)

    add_body(
        "El diagrama de clases muestra la arquitectura del sistema NutriPrecio, organizada en capas "
        "que separan la interfaz de usuario, la logica de negocio y el acceso a datos."
    )

    add_heading("5.2 Clases Principales", level=2)

    classes = [
        ("User:", "Extiende AbstractUser. Email unico, relacion con Store via owner."),
        ("Store:", "Tienda/pyme con nombre, slug, logo, sitio web, owner (FK a User)."),
        ("Product:", "Catalogo con nombre, marca, categoria, unidad, imagen, descripcion."),
        ("Category:", "Categorias jerarquicas con referencia padre-hijo."),
        ("Price:", "Precio vinculado a Product y Store. Incluye descuento y disponibilidad."),
        ("AuthService:", "Servicio Angular para login, registro, logout y gestion de token."),
        ("ApiService:", "Servicio Angular base para comunicaciones HTTP tipadas."),
        ("AuthGuard:", "Guard de rutas que verifica autenticacion y rol de vendedor."),
    ]
    for label, desc in classes:
        p = Paragraph(f'<b>{label}</b> {desc}', styles["ReportBody"])
        story.append(p)

    add_heading("5.3 Relaciones", level=2)

    for rel in [
        "User (1) -- (N) Store: Un vendedor puede tener una o mas tiendas.",
        "Category (1) -- (N) Product: Una categoria contiene multiples productos.",
        "Product (1) -- (N) Price: Un producto tiene multiples registros de precio.",
        "Store (1) -- (N) Price: Una tienda tiene multiples registros de precio.",
    ]:
        add_bullet(rel)

    # =====================
    # 7. DIAGRAMA DE ACTIVIDADES
    # =====================
    add_heading("6. Diagrama de Actividades")

    add_image_with_caption(
        os.path.join(DOCS_DIR, "diagrama-actividades.png"),
        "Figura 3: Diagrama de Actividades con los 3 flujos principales"
    )

    add_heading("6.1 Flujo MV-52: Login y Registro", level=2)

    add_body(
        "El usuario accede al formulario de login o registro. El sistema valida los campos, envia "
        "la solicitud al backend, que verifica credenciales contra la base de datos. Si son validas, "
        "genera un token de autenticacion y lo retorna al frontend. En registro, se crea el usuario "
        "con contrasena encriptada mediante create_user() de Django."
    )

    add_heading("6.2 Flujo MV-03: Registro de Tienda", level=2)

    add_body(
        "El vendedor autenticado accede al formulario desde el dashboard. El frontend construye "
        "FormData con los datos de la tienda y lo envia a POST /api/stores/. El backend valida el "
        "rol de vendedor, crea la instancia de Store vinculada al owner, y retorna la tienda creada."
    )

    add_heading("6.3 Flujo MV-54: Dashboard del Vendedor", level=2)

    add_body(
        "Al acceder a /dashboard, el authGuard verifica autenticacion y rol is_seller. Si ambas "
        "condiciones se cumplen, se carga la vista con mensaje de bienvenida personalizado, perfil "
        "del usuario con avatar, indicadores de estado, y acciones rapidas. Si no es vendedor, "
        "es redirigido al home."
    )

    # =====================
    # 8. SPRINT REVIEW
    # =====================
    add_heading("7. Sprint Review")

    add_heading("7.1 Funcionalidades Entregadas", level=2)

    delivered = [
        ("MV-52: Login y Registro (COMPLETO)",
         "Registro con encriptacion de contrasenas. Tokens de autenticacion. Interceptor HTTP "
         "authInterceptor. Guard authGuard con validacion de token y rol. Formularios con validacion."),
        ("MV-03: Registro de Tienda (COMPLETO)",
         "Modelo Store con owner FK. Serializer y ViewSet con permisos. Endpoint CRUD en /api/stores/. "
         "Formulario frontend store-form.component integrado al dashboard."),
        ("MV-54: Dashboard del Vendedor (COMPLETO)",
         "Mensaje de bienvenida personalizado. Perfil con avatar de iniciales. Indicadores de estado. "
         "Acciones rapidas con navegacion. Validacion de rol vendedor en authGuard."),
    ]
    for title, desc in delivered:
        p = Paragraph(f'<b><font color="#007000">{title}</font></b>', styles["ReportBody"])
        story.append(p)
        story.append(Paragraph(desc, styles["ReportBody"]))
        add_spacer(4)

    add_heading("7.2 Funcionalidades NO Entregadas / Parciales", level=2)

    partial = [
        ("Dashboard con sidebar:", "No se implemento el layout con barra lateral de navegacion."),
        ("Estado real del perfil:", "Sin integracion activa con API para datos reales de la tienda."),
        ("Gestion de productos:", "Boton deshabilitado con etiqueta Proximamente. Fuera del alcance del Sprint I."),
    ]
    for title, desc in partial:
        p = Paragraph(f'<b><font color="#BF8F00">{title}</font></b> {desc}', styles["ReportBody"])
        story.append(p)

    add_heading("7.3 Metricas del Sprint", level=2)

    metrics_headers = ["Indicador", "Valor"]
    metrics_data = [
        ["Tareas planificadas", "9"],
        ["Tareas completadas", "7 de 9 (78%)"],
        ["Tareas parciales", "2 de 9 (22%)"],
        ["Horas estimadas", "47 horas"],
        ["Horas consumidas", f"{TOTAL_CONSUMED} horas"],
    ]
    add_table(metrics_headers, metrics_data, col_widths=[5*cm, 8*cm])

    add_heading("7.4 Feedback Obtenido", level=2)

    add_body(
        "No se obtuvo feedback formal de stakeholders durante el Sprint I. Las revisiones se "
        "realizaron internamente entre miembros del equipo."
    )

    # =====================
    # 9. SPRINT RETROSPECTIVE
    # =====================
    add_heading("8. Sprint Retrospective")

    add_heading("8.1 Que Salio Bien", level=2)

    for item in [
        "Distribucion clara de tareas con responsabilidades bien definidas.",
        "Stack tecnologico apropiado: Django + Angular facilito desarrollo paralelo.",
        "Autenticacion implementada correctamente desde el inicio con tokens e interceptor.",
        "Estructura del proyecto organizada por dominio permitio trabajo sin conflictos.",
    ]:
        add_bullet(item)

    add_heading("8.2 Que Se Puede Mejorar", level=2)

    for item in [
        "Falta de standups diarios resulto en tareas estancadas sin deteccion temprana.",
        "Campo owner de Store no se incluyo desde el inicio, requiriendo migracion adicional.",
        "No se valido rol is_seller en el dashboard inicialmente.",
        "Sin datos de prueba hasta el final, imposibilitando demos intermedias.",
        "Formulario de registro de tienda subestimado en horas.",
    ]:
        add_bullet(item)

    add_heading("8.3 Acciones para el Proximo Sprint", level=2)

    for i, action in enumerate([
        "Implementar daily standups de 15 minutos.",
        "Crear seed data al inicio del Sprint.",
        "Revisar modelos contra diagramas UML antes de codificar.",
        "Definir criterios de aceptacion mas claros por tarea.",
        "Estimar con mas margen para tareas de integracion.",
    ], 1):
        add_bullet(f"{i}. {action}")

    # =====================
    # 10. CONCLUSIONES
    # =====================
    add_heading("9. Conclusiones")

    add_heading("9.1 Lecciones Aprendidas", level=2)

    for lesson in [
        "La comunicacion entre diseno UML e implementacion es fundamental para evitar retrabajo.",
        "Las tareas de integracion frontend-backend requieren mas tiempo del estimado.",
        "Los datos de prueba son esenciales desde el inicio del Sprint.",
        "La validacion de roles debe ser parte de los criterios de aceptacion desde el inicio.",
        "El stack Django + Angular es efectivo para desarrollo paralelo con API bien definida.",
    ]:
        add_bullet(lesson)

    add_heading("9.2 Plan para Sprint II", level=2)

    for item in [
        "Completar el layout del dashboard con sidebar de navegacion lateral.",
        "Integrar el dashboard con la API para datos reales de la tienda.",
        "Implementar la gestion de productos (CRUD) desde el dashboard.",
        "Desarrollar la funcionalidad de busqueda y comparacion de precios.",
        "Implementar daily standups y mejorar seguimiento de tareas.",
        "Crear y mantener seed data actualizada para pruebas y demos.",
    ]:
        add_bullet(item)

    add_spacer(12)
    add_body(
        "El Sprint I sento las bases tecnicas y organizacionales del proyecto NutriPrecio. A pesar "
        "de las tareas parciales, el equipo logro entregar un incremento funcional con autenticacion "
        "completa, registro de tiendas y dashboard del vendedor. Las lecciones aprendidas serviran "
        "para mejorar la planificacion y ejecucion del Sprint II."
    )

    # Build PDF
    doc.build(story)
    print(f"PDF saved: {output_path}")
    return output_path


if __name__ == "__main__":
    docx_path = create_docx()
    create_pdf(docx_path)
    print("Report generation complete.")
